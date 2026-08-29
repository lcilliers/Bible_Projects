"""prosestore.py — the DB-canonical prose store: extract, search, chapter export/import.

Incorporates operations that were previously standalone and untouched by any IBA code —
`scripts/build_programme_prose_extract.py`, `scripts/search_prose.py`,
`scripts/export_prose_chapter_edit.py`, `scripts/import_prose_chapter_edit.py` — into the app
proper (`governance.project_change_rule`: "Any operation defined in the past that is not in the
IBA app must be migrated to the app."). Escalation #784, 2026-08-21: the researcher's own finding
was that despite `prose_section`/`prose_section_type` being fully catalogued in `cfg_table`/
`cfg_column`, zero `cfg_step`/`cfg_work_package`/handler code ever touched them — the schema
catalogue described the tables, nothing in the app operated on them.

Connects to `bible_research.db` via `cfg.database_path('bible_research')` — the first real
operational consumer of that setting beyond `init.py`'s startup drift check (`Cfg.database_path`'s
own docstring notes both path settings "sat as genuine orphans" until now).

Config-driven values (module='prose' in `cfg_setting`) replace what were module-level hardcoded
constants flagged NON-COMPLIANT under escalation #648 (`iba/app/reports/hardcoded-constants-sweep-
20260817.md`): `prose.chapter_names`, `prose.book_stage_map`, `prose.search_default_limit`.
`EXTRACTOR_VERSION` (the fourth constant #648 flagged) is deliberately NOT one of these — researcher,
2026-08-22 (escalation #784/#794): it labelled the extract *tool's* own version, not prose content
(which has its own real per-section version/supersede model, `prose_section.version`/
`supersedes_id` — a completely different thing this constant was easily confused with). It
controlled nothing and gated nothing. Now that this file is under IBA's own code control (git,
`governance.build_md_on_code_change`), a static tool-version label has no reason to live in
`cfg_setting` too — dropped from the extract's output entirely rather than config-driven.

The four `scripts/*.py` CLI entry points remain the documented CLI usage
(`docs/prose-store-architecture.md` §8) — they now import their core logic from here instead of
defining it locally, so there is exactly one implementation, exercised both by direct CLI use and
by the registered `prose` work package (`python -m iba.app.run prose --step prose.extract ...`).

`run_flag()` (escalation #829 sec12.4, angle a) is the one exception to "read-only against
`prose_section`/`prose_section_type`" above — it writes directly to `wa_data_quality_flags`
(`bible_research.db`), a different table family entirely, not gated behind a generated patch file.

`run_flag_fix_propose()`/`run_flag_fix_apply()` (escalation #890 D5, angle b) add propose/apply on
top of angle a: propose writes a review report (no DB write); apply re-checks each approved
section fresh and generates a `PROSE` supersede patch, same shape/boundary as `run_import_chapter`
— neither is a new exception to the read-only/patch-generating rule above. `run_import_chapter`
also gained delete-detection (escalation #890 D3) — a section missing from an edit file now
refuses the import, matching add/move's existing behaviour, instead of silently no-op'ing.
"""
from __future__ import annotations

import json
import re
import sqlite3
from datetime import datetime, timezone
from pathlib import Path

# Every output-location constant that used to live here (OUT_DIR/DOCX_OUT_DIR/SEARCH_OUT_DIR/
# CHAPTER_EDIT_OUT_DIR/PATCH_OUT_DIR/_DEFAULT_CHAPTER_NAMES/_DEFAULT_BOOK_STAGE_MAP/
# _DEFAULT_BOOK_OUTPUT_DIR/_DEFAULT_SEARCH_LIMIT/_DEFAULT_EDIT_FILE_DIR) was a "Python-level
# DEFAULT only, used when cfg_prose is inactive/absent" — removed entirely 2026-08-29 (researcher
# direct ruling, full-codebase sweep: *"there should be NO hardcoded literals in current code"*).
# This was, on its own history, the clearest case for removing the pattern rather than re-fixing
# its drifted value again: CHAPTER_EDIT_OUT_DIR had ALREADY silently drifted once (corrected
# 2026-08-22, drifted again, corrected again 2026-08-28, escalation #971/#976) — the fallback
# constant, not the config, was what kept going stale, because nothing ever re-verified it against
# the live row. Every function below now reads its `cfg_prose` row via `cfg.required_module_
# setting()` (lib/cfg.py, no default parameter at all) — a missing/inactive row now fails loudly
# instead of silently resurrecting a value that already proved it can go stale unnoticed.

MARKER_RE = re.compile(r"<!-- PROSE_([A-Z_]+): ?(.*?) -->")
ID_RE = re.compile(r"<!-- PROSE_SECTION_ID: (\d+) -->")


def open_db(cfg) -> sqlite3.Connection:
    """The one connection point — `bible_research.db` located via config, not a literal path."""
    conn = sqlite3.connect(cfg.database_path("bible_research"))
    conn.row_factory = sqlite3.Row
    return conn


def now_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def today_compact() -> str:
    return datetime.now(timezone.utc).strftime("%Y%m%d")


def _next_edit_version(stem: str, edit_dir: Path) -> int:
    """Next edit-cycle version number for a chapter-edit export with this book/chapter/section
    stem. `edit_dir` has no default — its sole caller always passes `edit_file_dir(cfg)` explicitly
    (a hardcoded default here would be the exact same drift-prone shape removed everywhere else
    2026-08-29). Researcher, 2026-08-22 (escalation #784): 'all files must be version controlled' --
    'currently the name of the file makes it impossible to link the file with the book-chapter-
    session' (read as: without a version number, two edit exports of the same book+chapter can't be
    told apart, and a same-day re-export silently overwrites). This is an edit-CYCLE version, distinct from
    prose_section.version -- one exported file can bundle several section rows that may each sit at
    a different DB version, so the file's own version is a separate counter. Scans both the active
    folder and its archive (§5 below moves imported files there) so a version is never reused even
    after archiving."""
    pattern = re.compile(re.escape(stem) + r"-v(\d+)-\d{8}\.md$")
    max_v = 0
    for folder in (edit_dir, edit_dir / "archive"):
        if not folder.exists():
            continue
        for f in folder.iterdir():
            m = pattern.match(f.name)
            if m:
                max_v = max(max_v, int(m.group(1)))
    return max_v + 1


def get_schema_version(conn) -> str:
    row = conn.execute(
        "SELECT version_code FROM schema_version WHERE id = (SELECT MAX(id) FROM schema_version)"
    ).fetchone()
    return row[0] if row else "unknown"


def chapter_names(cfg) -> dict:
    """`prose.chapter_names` — keys are strings (JSON object keys always are); callers look up by
    `str(chapter_no)`. Reads the dedicated `cfg_prose` module table (escalation #829,
    `governance.module.config`), not generic `cfg_setting` — corrects this function's own earlier
    (pre-#829) use of `cfg.setting()`, which duplicated table-driven module config into the
    project-wide settings table."""
    return cfg.required_module_setting("cfg_prose", "prose.chapter_names")


def book_stage_map(cfg) -> dict:
    """`prose.book_stage_map` — D10 RESOLVED (escalation #890 D6, 2026-08-26): this function's own
    docstring previously claimed the stage-based map drives which book a `prose_section_type` row
    is filed under, and that 1 of 949 rows (id 78, `prog_purp_observations_framework` —
    `source_stage='programme'`, `book_label='Detail design'`) would be misfiled as a result. That
    claim was checked against the actual call sites this round and found **false** — this function
    is used ONLY to validate the `--book` CLI argument against the list of real book names
    (`run_extract`'s `book not in book_stage_map(cfg)` check); the actual row filtering
    (`extract_programme_prose`) already queries `WHERE book_label = ?` directly and always has.
    Id 78 is therefore already correctly filed under 'Detail design' (its own `book_label`), not
    'Programme' — there was no functional bug, only a stale comment describing a design that was
    superseded by the time the code was actually written. Kept as the `--book` choice-list source
    (still useful for CLI validation) — not dropped, since nothing else currently enumerates the 4
    live book names."""
    return cfg.required_module_setting("cfg_prose", "prose.book_stage_map")


def search_default_limit(cfg) -> int:
    return cfg.required_module_setting("cfg_prose", "prose.search_default_limit")


def edit_file_dir(cfg) -> Path:
    """`prose.edit_file_dir` (escalation #829, `governance.rules_must_be_config_driven`) — no
    hardcoded fallback (removed 2026-08-29; its predecessor constant drifted silently twice)."""
    return Path(cfg.required_module_setting("cfg_prose", "prose.edit_file_dir"))


def output_dir(cfg) -> Path:
    """`prose.output_dir` (escalation #971/#976, `iba/app/lib/pathaudit.py`'s scan) — no hardcoded
    fallback (removed 2026-08-29, same reasoning as `edit_file_dir` above). Stays the Programme-
    only/no-`--book`-given default — `output_dir_for` below is the book-aware entry point
    (escalation #989/#1000)."""
    return Path(cfg.required_module_setting("cfg_prose", "prose.output_dir"))


def output_dir_for(cfg, book_label: str | None) -> Path:
    """`prose.book_output_dir` — book-aware output directory (escalation #989/#1000). These
    folders hold prose operation WORKING FILES for that book, not a replica of prose content
    (researcher ruling, #1000). `book_label is None` (no `--book` given) falls back to the flat
    `output_dir(cfg)` (Programme's own dir), unchanged existing behaviour. A `book_label` not
    present in the map raises rather than guessing a location -- matches `run_extract`'s existing
    "unknown book" refusal for an unrecognised `--book` value."""
    if book_label is None:
        return output_dir(cfg)
    book_map = cfg.required_module_setting("cfg_prose", "prose.book_output_dir")
    if book_label not in book_map:
        raise ValueError(
            f"no prose.book_output_dir entry for book {book_label!r}; choose from: "
            f"{', '.join(book_map)}")
    return Path(book_map[book_label])


def docx_output_dir(cfg) -> Path:
    """`prose.docx_output_dir` — same reasoning as `output_dir` above."""
    return Path(cfg.required_module_setting("cfg_prose", "prose.docx_output_dir"))


def search_output_dir(cfg) -> Path:
    """`prose.search_output_dir` — same reasoning as `output_dir` above."""
    return Path(cfg.required_module_setting("cfg_prose", "prose.search_output_dir"))


def patch_output_dir(cfg) -> Path:
    """`prose.patch_output_dir` — same reasoning as `output_dir` above."""
    return Path(cfg.required_module_setting("cfg_prose", "prose.patch_output_dir"))


# ── extract (was scripts/build_programme_prose_extract.py) ─────────────────

def extract_programme_prose(
    conn, include_body: bool = False, book: str | None = None, chapter: int | None = None
) -> dict:
    book_filter = ""
    params: tuple = ()
    if book is not None:
        book_filter = " AND book_label = ?"
        params = (book,)
    if chapter is not None:
        book_filter += " AND chapter_no = ?"
        params += (chapter,)
    type_rows = conn.execute(
        f"""SELECT id, code, label, description, chapter_no, lifecycle_tag,
                  expected_length_min, expected_length_max, sort_order,
                  book_order, book_label, section_order, section_label,
                  source_stage
             FROM prose_section_type
            WHERE delete_flagged = 0{book_filter}
            ORDER BY book_order, section_order, chapter_no, sort_order, id, code""",
        params,
    ).fetchall()

    types: list[dict] = []
    type_count = 0
    section_total = 0
    for t in type_rows:
        type_count += 1
        sections = conn.execute(
            """SELECT id, registry_id, heading, status, version, author,
                      word_count, created_at, approved_at
                 FROM prose_section
                WHERE section_type_id = ?
                  AND delete_flagged = 0
                ORDER BY version DESC""",
            (t["id"],),
        ).fetchall()
        entry = {
            "id": t["id"],
            "code": t["code"],
            "label": t["label"],
            "description": t["description"],
            "chapter_no": t["chapter_no"],
            "lifecycle_tag": t["lifecycle_tag"],
            "expected_length_min": t["expected_length_min"],
            "expected_length_max": t["expected_length_max"],
            "sort_order": t["sort_order"],
            "book_order": t["book_order"],
            "book_label": t["book_label"],
            "section_order": t["section_order"],
            "section_label": t["section_label"],
            "source_stage": t["source_stage"],
            "section_count": len(sections),
            "sections_preview": [dict(s) for s in sections],
        }
        if include_body and sections:
            bodies = {}
            for s in sections:
                body_row = conn.execute(
                    "SELECT body FROM prose_section WHERE id = ?", (s["id"],)
                ).fetchone()
                bodies[s["id"]] = body_row["body"] if body_row else None
            entry["bodies_by_id"] = bodies
        types.append(entry)
        section_total += len(sections)

    return {
        "types": types,
        "type_count": type_count,
        "section_total": section_total,
        "book": book,
        "chapter": chapter,
    }


def build_extract(
    conn, include_body: bool = False, book: str | None = None, chapter: int | None = None
) -> dict:
    pp = extract_programme_prose(conn, include_body=include_body, book=book, chapter=chapter)
    return {
        "meta": {
            "generated_at": now_iso(),
            "schema_version": get_schema_version(conn),
            "source": "prose_section_type + prose_section",
            "canonical_note": "DB is source of truth post-M34 for programme-stage narrative (L2 of 3-layer reference).",
            "include_body": include_body,
            "description": "Programme-wide narrative section index — anchor verse definition, XREF architecture, validation standard, etc. Content populated via PROSE patches as researcher + AI draft each narrative.",
            "patch_hint": "To insert programme-wide prose via a PROSE patch, use `section_type_id = <id>` from each entry in programme_prose.types. Pair with `registry_id = null` (requires the schema enablement directive per wa-directive-instruction [current] §10 to relax the NOT NULL constraint).",
        },
        "programme_prose": pp,
    }


def render_markdown_view(cfg, extract: dict) -> str:
    names = chapter_names(cfg)
    lines = []
    meta = extract["meta"]
    pp = extract["programme_prose"]

    book_name = pp.get("book") or "All books"
    chapter_hdr = f" — Chapter {pp['chapter']}" if pp.get("chapter") is not None else ""
    lines.append(f"# Prose Extract — {book_name}{chapter_hdr} — {meta['generated_at'][:10]}\n")
    lines.append(f"_Schema {meta['schema_version']} · source: `prose_section_type` + `prose_section`._\n")
    lines.append("---\n")
    lines.append("## Summary\n")
    lines.append(f"**Section types seeded:** {pp['type_count']}  ·  "
                 f"**Content sections populated:** {pp['section_total']}\n")
    lines.append("---\n")

    populated = [t for t in pp["types"] if t["section_count"] > 0]
    stubs = [t for t in pp["types"] if t["section_count"] == 0]

    if populated:
        by_book: dict = {}
        for t in populated:
            by_book.setdefault((t.get("book_order"), t.get("book_label") or "Unassigned book"), []).append(t)

        book_keys = sorted(by_book.keys(), key=lambda x: (x[0] is None, x[0] if x[0] is not None else 0, x[1]))
        for book_key in book_keys:
            lines.append(f"## {book_key[1]}\n")
            by_section: dict = {}
            for t in by_book[book_key]:
                section_key = (t.get("section_order"), t.get("section_label") or t.get("source_stage") or "Unassigned section")
                by_section.setdefault(section_key, []).append(t)
            section_keys = sorted(by_section.keys(), key=lambda x: (x[0] is None, x[0] if x[0] is not None else 0, x[1]))
            for section_key in section_keys:
                if section_key[1] != book_key[1]:
                    lines.append(f"### {section_key[1]}\n")
                by_chapter: dict = {}
                for t in by_section[section_key]:
                    by_chapter.setdefault(t.get("chapter_no"), []).append(t)
                chapter_keys = sorted(by_chapter.keys(), key=lambda c: (c is None, c if c is not None else 0))
                for ch in chapter_keys:
                    chapter_name = names.get(str(ch), "Unchaptered") if ch is not None else "Unchaptered"
                    header = f"Chapter {ch} — {chapter_name}" if ch is not None else chapter_name
                    lines.append(f"#### {header}\n")
                    for t in sorted(by_chapter[ch], key=lambda x: (x["sort_order"] is None, x["sort_order"] or 0, x["id"])):
                        if t.get("description"):
                            lines.append(f"> {t['description']}\n")
                        lines.append(f"##### {t['label']}\n")
                        lines.append(f"_`{t['code']}`  ·  type id {t['id']}  ·  chapter {t['chapter_no']}  ·  sort {t['sort_order']}_\n")
                        bodies = t.get("bodies_by_id") or {}
                        for s in t["sections_preview"]:
                            body = bodies.get(s["id"])
                            if body is None:
                                lines.append(f"*(metadata only — body not included in this extract. "
                                             f"Run `build_programme_prose_extract.py --include-body` to render full text.)*")
                                lines.append(f"- Section id {s['id']} · status `{s['status']}` · "
                                             f"v{s['version']} · {s['word_count']} words · author `{s['author']}`\n")
                            else:
                                meta_line = (f"*Section id {s['id']} · status `{s['status']}` · "
                                             f"v{s['version']} · {s['word_count']} words · author `{s['author']}`*")
                                lines.append(meta_line + "\n")
                                lines.append(body.rstrip())
                                lines.append("")
                        lines.append("")

    if stubs:
        lines.append("---\n")
        lines.append("## Section types not yet populated\n")
        lines.append("_Stubs — `prose_section_type` rows with `chapter_no=NULL` or no `prose_section` content. "
                     "`id` is the value to use as `section_type_id` in a PROSE patch insert._\n")
        lines.append("| id | code | label | chapter | sort | description |")
        lines.append("|---:|---|---|---:|---:|---|")
        for t in sorted(stubs, key=lambda x: (x["chapter_no"] is None, x["chapter_no"] or 0, x["sort_order"] or 0)):
            desc = (t.get("description") or "").replace("|", "\\|").replace("\n", " ")
            lines.append(f"| {t['id']} | `{t['code']}` | {t['label']} | "
                         f"{t['chapter_no'] if t['chapter_no'] is not None else '—'} | "
                         f"{t['sort_order']} | {desc} |")
        lines.append("")

    lines.append("---")
    lines.append(f"*Generated {meta['generated_at']}.*")
    return "\n".join(lines)


def render_docx(cfg, extract: dict, out_path: Path) -> bool:
    """Readable .docx export. Returns False (no-op) if python-docx isn't installed."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return False

    names = chapter_names(cfg)
    meta = extract["meta"]
    pp = extract["programme_prose"]

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Programme Prose", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle = doc.add_paragraph()
    r = subtitle.add_run(
        f"Schema {meta['schema_version']} · generated {meta['generated_at']}"
    )
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    summary = doc.add_paragraph()
    r = summary.add_run(
        f"Section types seeded: {pp['type_count']}  ·  "
        f"Content sections populated: {pp['section_total']}"
    )
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    populated = [t for t in pp["types"] if t["section_count"] > 0]
    stubs = [t for t in pp["types"] if t["section_count"] == 0]

    if populated:
        doc.add_heading("Programme", level=1)
        by_chapter: dict = {}
        for t in populated:
            by_chapter.setdefault(t.get("chapter_no"), []).append(t)
        chapter_keys = sorted(by_chapter.keys(), key=lambda c: (c is None, c if c is not None else 0))

        for ch in chapter_keys:
            chapter_name = names.get(str(ch), "Unchaptered") if ch is not None else "Unchaptered"
            header = f"Chapter {ch} — {chapter_name}" if ch is not None else chapter_name
            doc.add_heading(header, level=2)

            for t in sorted(by_chapter[ch], key=lambda x: (x["sort_order"] or 0)):
                doc.add_heading(t["label"], level=3)
                meta_p = doc.add_paragraph()
                rr = meta_p.add_run(f"code: {t['code']}  ·  type id {t['id']}  ·  sort {t['sort_order']}")
                rr.italic = True
                rr.font.size = Pt(9)
                rr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

                if t.get("description"):
                    desc_p = doc.add_paragraph(style="Intense Quote")
                    desc_p.add_run(t["description"])

                bodies = t.get("bodies_by_id") or {}
                for s in t["sections_preview"]:
                    sec_meta = doc.add_paragraph()
                    rr = sec_meta.add_run(
                        f"Section id {s['id']}  ·  status {s['status']}  ·  v{s['version']}  ·  "
                        f"{s['word_count']} words  ·  author {s['author']}"
                    )
                    rr.italic = True
                    rr.font.size = Pt(9)
                    rr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

                    body = bodies.get(s["id"])
                    if body is None:
                        p = doc.add_paragraph()
                        rr = p.add_run("(body not included in this extract — re-run with --include-body)")
                        rr.italic = True
                        continue
                    for para in body.split("\n\n"):
                        para = para.strip()
                        if not para:
                            continue
                        doc.add_paragraph(para)

    if stubs:
        doc.add_heading("Section types not yet populated", level=1)
        caption = doc.add_paragraph()
        rr = caption.add_run(
            "Stubs — prose_section_type rows with chapter_no=NULL or no prose_section content. "
            "id is the value to use as section_type_id in a PROSE patch insert."
        )
        rr.italic = True
        rr.font.size = Pt(9)
        rr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = (
            "id", "code", "label", "chapter", "description")
        for t in sorted(stubs, key=lambda x: (x["chapter_no"] is None, x["chapter_no"] or 0, x["sort_order"] or 0)):
            row = table.add_row().cells
            row[0].text = str(t["id"])
            row[1].text = t["code"]
            row[2].text = t["label"]
            row[3].text = str(t["chapter_no"]) if t["chapter_no"] is not None else "—"
            row[4].text = (t.get("description") or "").replace("\n", " ")

    doc.save(out_path)
    return True


def run_extract(cfg, include_body=False, book=None, chapter=None, also_markdown=False,
                also_docx=False, out=None) -> dict:
    """High-level orchestration shared by the CLI script and the `prose.extract` handler.
    Returns {"json": path, "md": path|None, "docx": path|None, "type_count": n, "section_total": n}."""
    conn = open_db(cfg)
    try:
        if book is not None and book not in book_stage_map(cfg):
            raise ValueError(f"unknown book {book!r}; choose from: {', '.join(book_stage_map(cfg))}")
        if chapter is not None and book is None:
            raise ValueError("chapter requires book")

        extract = build_extract(conn, include_body=include_body, book=book, chapter=chapter)

        out_dir = output_dir_for(cfg, book)
        out_dir.mkdir(parents=True, exist_ok=True)
        stamp = today_compact()
        out_path = Path(out) if out else out_dir / f"wa-programme-prose-extract-{stamp}.json"
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(json.dumps(extract, indent=2, ensure_ascii=False), encoding="utf-8")

        result = {"json": str(out_path), "md": None, "docx": None,
                  "type_count": extract["programme_prose"]["type_count"],
                  "section_total": extract["programme_prose"]["section_total"]}

        if also_markdown:
            md_path = out_path.with_suffix(".md")
            md_path.write_text(render_markdown_view(cfg, extract), encoding="utf-8")
            result["md"] = str(md_path)

        if also_docx:
            docx_dir = docx_output_dir(cfg)
            docx_dir.mkdir(parents=True, exist_ok=True)
            docx_path = docx_dir / f"wa-programme-prose-extract-{stamp}.docx"
            extract_wb = extract if include_body else build_extract(conn, include_body=True)
            wrote = render_docx(cfg, extract_wb, docx_path)
            result["docx"] = str(docx_path) if wrote else None

        return result
    finally:
        conn.close()


# ── search (was scripts/search_prose.py) ────────────────────────────────────

def search_prose(conn, query, book=None, limit=100, raw_fts=False):
    book_clause = ""
    fts_query = query if raw_fts else f'"{query.replace(chr(34), chr(34) + chr(34))}"'
    params: list = [fts_query]
    if book is not None:
        book_clause = "AND lower(pst.book_label) = lower(?)"
        params.append(book.strip())
    params.append(limit)

    where_sql = f"""
        FROM prose_section_fts
        JOIN prose_section ps ON ps.id = prose_section_fts.rowid
        JOIN prose_section_type pst ON pst.id = ps.section_type_id
        WHERE prose_section_fts MATCH ?
          AND COALESCE(ps.delete_flagged, 0) = 0
          AND COALESCE(pst.delete_flagged, 0) = 0
          {book_clause}
    """
    total_params = params[:-1]
    total_matches = conn.execute(f"SELECT COUNT(*) {where_sql}", total_params).fetchone()[0]

    rows = conn.execute(
        f"""
        SELECT
            ps.id, ps.heading, ps.status, ps.version, ps.author,
            pst.source_stage, pst.book_order, pst.book_label, pst.section_order,
            pst.section_label, pst.chapter_no, pst.label AS chapter_title, pst.sort_order,
            snippet(prose_section_fts, 0, '**', '**', ' ... ', 36) AS context,
            bm25(prose_section_fts) AS relevance
        {where_sql}
        ORDER BY relevance, pst.book_order, pst.section_order,
                 pst.chapter_no, pst.sort_order, ps.id
        LIMIT ?
        """,
        params,
    ).fetchall()
    return rows, total_matches


def render_search_markdown(query, rows, total_matches, book):
    generated = now_iso()
    scope = book or "All books"
    lines = [
        f"# Prose Search — {query}", "",
        f"_Generated {generated} · scope: {scope} · showing {len(rows)} of {total_matches} matches_",
        "", "---", "",
    ]
    if not rows:
        lines.append("No active prose sections matched the query.")
        return "\n".join(lines) + "\n"

    for index, row in enumerate(rows, start=1):
        location = [row["book_label"] or "Unassigned book",
                   row["section_label"] or row["source_stage"] or "Unassigned section"]
        if row["chapter_no"] is not None:
            location.append(f"Chapter {row['chapter_no']}: {row['chapter_title']}")
        lines.extend([
            f"## {index}. {row['heading']}", "",
            f"**Reference:** {' / '.join(location)}",
            f"**Prose section ID:** `{row['id']}` · **status:** `{row['status']}` · **version:** `{row['version']}`",
            "",
            f"> {row['context']}", "",
        ])
    return "\n".join(lines)


def default_search_output_path(cfg, query, book):
    slug = re.sub(r"[^a-z0-9]+", "-", query.lower()).strip("-") or "query"
    if book:
        book_slug = re.sub(r"[^a-z0-9]+", "-", book.lower()).strip("-")
        slug = f"{slug}-{book_slug}"
    return search_output_dir(cfg) / f"prose-search-{slug}-{today_compact()}.md"


def run_search(cfg, query, book=None, limit=None, raw_fts=False, out=None) -> dict:
    conn = open_db(cfg)
    try:
        eff_limit = limit if limit is not None else search_default_limit(cfg)
        rows, total_matches = search_prose(conn, query, book=book, limit=eff_limit, raw_fts=raw_fts)
        report = render_search_markdown(query, rows, total_matches, book)
        output_path = Path(out) if out else default_search_output_path(cfg, query, book)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(report, encoding="utf-8")
        return {"path": str(output_path), "shown": len(rows), "total": total_matches}
    finally:
        conn.close()


# ── chapter export (was scripts/export_prose_chapter_edit.py) ──────────────

def run_export_chapter(cfg, type_id=None, book=None, chapter=None, out=None) -> dict:
    conn = open_db(cfg)
    try:
        if type_id is not None:
            sql = """
            SELECT ps.id, ps.heading, ps.body, ps.version,
                   pst.code, pst.book_order, pst.book_label, pst.section_order,
                   pst.section_label, pst.chapter_no, pst.label AS chapter_title,
                   pst.description, pst.sort_order
            FROM prose_section ps
            JOIN prose_section_type pst ON pst.id = ps.section_type_id
            WHERE pst.id = ?
              AND COALESCE(pst.delete_flagged, 0) = 0
              AND COALESCE(ps.delete_flagged, 0) = 0
            ORDER BY pst.sort_order, ps.id
            """
            query_params = (type_id,)
        else:
            if chapter is None:
                raise ValueError("chapter is required with book")
            sql = """
            SELECT ps.id, ps.heading, ps.body, ps.version,
                   pst.code, pst.book_order, pst.book_label, pst.section_order,
                   pst.section_label, pst.chapter_no, pst.label AS chapter_title,
                   pst.description, pst.sort_order
            FROM prose_section ps
            JOIN prose_section_type pst ON pst.id = ps.section_type_id
            WHERE lower(pst.book_label) = lower(?)
              AND pst.chapter_no = ?
              AND COALESCE(pst.delete_flagged, 0) = 0
              AND COALESCE(ps.delete_flagged, 0) = 0
            ORDER BY pst.sort_order, ps.id
            """
            query_params = (book, chapter)
        rows = conn.execute(sql, query_params).fetchall()

        if not rows:
            raise ValueError(f"no active prose rows found for {book!r}, chapter {chapter}")

        stamp = today_compact()
        book_name = rows[0]["book_label"] or "unassigned-book"
        chapter_no = rows[0]["chapter_no"]
        title = f"{book_name} — Chapter {chapter_no}" if chapter_no is not None else book_name
        # Filename stem: book + chapter, or book + section code when there's no chapter to key on
        # (e.g. a single-type export). Researcher, 2026-08-22 (escalation #784): the un-versioned
        # name made it "impossible to link the file with the book-chapter-session" -- two exports
        # of the same book/chapter were indistinguishable, and a same-day re-export silently
        # overwrote (found live this session). -v{n}- (an edit-cycle counter, see
        # _next_edit_version) is what makes each export traceable to its own edit session.
        stem = (
            f"prose-edit-{book_name.lower().replace(' ', '-')}-"
            f"{('chapter-' + str(chapter_no)) if chapter_no is not None else ('section-' + str(rows[0]['code']))}"
        )
        edit_dir = edit_file_dir(cfg)
        output = Path(out) if out else edit_dir / (
            f"{stem}-v{_next_edit_version(stem, edit_dir)}-{stamp}.md"
        )
        lines = [
            f"# Prose Edit — {title}", "",
            "<!-- Edit only the prose body below each chapter heading. Do not change markers. -->",
            "<!-- This file becomes permanent provenance once imported (its archived path is -->",
            "<!-- recorded as record_change_log.change_source, escalation #836) -- do not delete -->",
            "<!-- by hand; the import step archives it automatically on success. -->",
            # Escalation #890 D3: the authoritative record of which section ids this export
            # covered -- read back at import time to detect a whole block silently removed by
            # hand. Deliberately NOT re-derived from whichever blocks happen to survive in the
            # file (a section_type-code-based inference was tried and found live to blind itself
            # exactly when the deleted block was the ONLY one of its type left in the file --
            # its code then vanishes from the surviving blocks too, so nothing is left to detect
            # against). This marker is fixed at export time and can't be fooled by any deletion.
            f"<!-- PROSE_EXPORT_SECTION_IDS: {','.join(str(r['id']) for r in rows)} -->", "",
        ]
        for row in rows:
            lines.extend([
                f"<!-- PROSE_SECTION_ID: {row['id']} -->",
                f"<!-- PROSE_SECTION_TYPE: {row['code']} -->",
                f"<!-- PROSE_BOOK: {row['book_label']} -->",
                f"<!-- PROSE_SECTION: {row['section_label']} -->",
                f"<!-- PROSE_CHAPTER_NO: {row['chapter_no']} -->",
                f"<!-- PROSE_CHAPTER_TITLE: {row['chapter_title']} -->",
                f"<!-- PROSE_SORT_ORDER: {row['sort_order']} -->",
                f"<!-- PROSE_VERSION: {row['version']} -->",
                "", f"## {row['heading']}", "", row["body"].rstrip(), "", "---", "",
            ])

        output.parent.mkdir(parents=True, exist_ok=True)
        output.write_text("\n".join(lines), encoding="utf-8")
        return {"path": str(output), "sections": len(rows)}
    finally:
        conn.close()


# ── chapter import (was scripts/import_prose_chapter_edit.py) ──────────────

def _parse_edit_blocks(text: str) -> list[dict[str, str]]:
    starts = list(ID_RE.finditer(text))
    if not starts:
        raise ValueError("no PROSE_SECTION_ID markers found")
    blocks = []
    for index, match in enumerate(starts):
        end = starts[index + 1].start() if index + 1 < len(starts) else len(text)
        chunk = text[match.start():end]
        markers = {key: value.strip() for key, value in MARKER_RE.findall(chunk)}
        heading_match = re.search(r"^##\s+(.+?)\s*$", chunk, re.MULTILINE)
        if not heading_match:
            raise ValueError(f"section {markers.get('SECTION_ID')} has no ## heading")
        body = chunk[heading_match.end():].strip()
        body = re.sub(r"\n---\s*$", "", body).strip()
        markers["BODY"] = body
        blocks.append(markers)
    return blocks


def run_import_chapter(cfg, input_path, author="researcher", out=None) -> dict:
    """Validates an edited chapter file and generates a PROSE supersede patch. Never writes to the
    database itself — apply the reviewed patch with scripts/apply_session_patch.py, same as before
    this operation was incorporated (the write-authorisation boundary is unchanged)."""
    input_path = Path(input_path)
    # Computed up front (before any move happens) so the patch's own _patch_summary can name it
    # even though the physical move only happens after the patch is successfully written.
    archived_source = edit_file_dir(cfg) / "archive" / input_path.name
    text = input_path.read_text(encoding="utf-8")
    blocks = _parse_edit_blocks(text)
    conn = open_db(cfg)
    operations = []
    try:
        # Delete-detection (escalation #890 D3, resolving #784 sec6's open decision): a section
        # silently vanishing from an edit file used to be a no-op -- the removed row's DB state
        # came back completely untouched, no error, no trace. Add and move both already refuse
        # outright (found live at #784); this makes delete symmetric with them, for the same
        # reason -- an edit file is a round-trip artefact (export -> edit -> import), not an
        # authoring surface for structural changes. Compares against the PROSE_EXPORT_SECTION_IDS
        # marker `run_export_chapter` now writes -- a fixed record of the export's own original
        # scope, not re-derived from whichever blocks happen to survive (an earlier
        # section_type-code-based version of this check was tried and found live, by actual
        # testing, to blind itself exactly when the deleted block was the only one of its type
        # left in the file). Files exported before this fix carry no marker -- skipped, not
        # crashed, with a visible note, rather than refusing every pre-existing edit file in
        # flight.
        export_ids_match = re.search(r"<!-- PROSE_EXPORT_SECTION_IDS: ([\d,]*) -->", text)
        if export_ids_match:
            expected_ids = {int(x) for x in export_ids_match.group(1).split(",") if x}
            seen_ids = {int(b["SECTION_ID"]) for b in blocks if b.get("SECTION_ID")}
            missing_ids = expected_ids - seen_ids
            if missing_ids:
                raise ValueError(
                    f"{len(missing_ids)} section(s) {sorted(missing_ids)} present in this "
                    f"file's original export are missing from it now -- refusing to import "
                    f"(escalation #890 D3: a section vanishing from an edit file is refused, "
                    f"matching add/move's existing behaviour, not silently ignored). If "
                    f"retiring a section is genuinely intended, do that explicitly "
                    f"(status='archived') rather than by omission from an edit file.")
        else:
            print(f"  [NOTE] {input_path} has no PROSE_EXPORT_SECTION_IDS marker (exported "
                  f"before escalation #890's delete-detection fix) -- delete-detection skipped "
                  f"for this file.")
        for block in blocks:
            required = ["SECTION_ID", "SECTION_TYPE", "BOOK", "SECTION", "CHAPTER_NO",
                       "CHAPTER_TITLE", "SORT_ORDER", "VERSION", "BODY"]
            missing = [key for key in required if key not in block]
            if missing:
                raise ValueError(f"section {block.get('SECTION_ID')} missing markers: {missing}")
            row = conn.execute(
                """
                SELECT ps.id, ps.version, ps.heading, ps.body,
                       pst.code, pst.book_label, pst.section_label,
                       pst.chapter_no, pst.label AS chapter_title, pst.sort_order
                FROM prose_section ps
                JOIN prose_section_type pst ON pst.id = ps.section_type_id
                WHERE ps.id = ? AND COALESCE(ps.delete_flagged, 0) = 0
                """,
                (int(block["SECTION_ID"]),),
            ).fetchone()
            if not row:
                raise ValueError(f"section {block['SECTION_ID']} is not an active current prose row")
            checks = {
                "SECTION_TYPE": row["code"], "BOOK": row["book_label"], "SECTION": row["section_label"],
                "CHAPTER_NO": str(row["chapter_no"]), "CHAPTER_TITLE": row["chapter_title"],
                "SORT_ORDER": str(row["sort_order"]), "VERSION": str(row["version"]),
            }
            for key, expected in checks.items():
                if block[key] != (expected or ""):
                    raise ValueError(
                        f"section {row['id']} marker {key} changed: "
                        f"file={block[key]!r}, database={expected!r}")
            if not block["BODY"]:
                raise ValueError(f"section {row['id']} has an empty body")
            # The section is the editing unit, not the chapter a bundled export happens to cover
            # (researcher, 2026-08-22, escalation #784: "each chapter can have multiple sections,
            # the section is the editing unit... if a chapter export is exported then all the
            # sections in the chapter will have to change version which is not necessary"). A
            # chapter-edit file bundles several independently-versioned sections purely for editing
            # convenience -- importing it back must only supersede the sections whose body actually
            # changed, not every block the file happens to contain.
            if block["BODY"].strip() == (row["body"] or "").strip():
                continue
            operations.append({
                "op_id": f"PROSE-{row['id']}-SUPERSEDE",
                "table": "prose_section",
                "operation": "supersede",
                # `supersedes_id` names the row being edited (Model A, escalation #836 --
                # apply_session_patch.py mutates it in place, no new row is created).
                "supersedes_id": row["id"],
                "record": {
                    "body": block["BODY"], "heading": row["heading"], "author": author,
                    "status": "draft",
                    "metadata_json": json.dumps({
                        "roundtrip_import": "iba.app.lib.prosestore.run_import_chapter",
                        "source_version": row["version"],
                    }),
                },
            })
    finally:
        conn.close()

    if not operations:
        raise ValueError(
            f"no changed sections in {input_path} -- every section's body matches the current "
            f"database row, nothing to import. The file is left in place (not archived): an "
            f"unedited export is still a disposable draft, not provenance for anything.")

    stamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    patch_id = f"PATCH-{stamp}-PROSE-CHAPTER-SUPERSEDE"
    patch = {
        "_patch_meta": {
            "patch_id": patch_id, "patch_type": "PROSE", "produced_at": now_iso(),
            "session_b_status": None, "researcher_approval": "PENDING",
            "description": f"Supersede {len(operations)} prose section(s) from an edited chapter Markdown file.",
        },
        "operations": operations,
        "_patch_summary": {
            "total_operations": len(operations), "prose_section_supersedes": len(operations),
            "source_edit_file": str(archived_source).replace("\\", "/"),
        },
    }
    output = Path(out) if out else patch_output_dir(cfg) / f"wa-prose-chapter-supersede-{stamp}.json"
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(patch, indent=2, ensure_ascii=False), encoding="utf-8")
    # On successful patch generation, archive the edit file -- researcher, 2026-08-22 (escalation
    # #784): "the import must get the file from the editing location, and on successful update move
    # the file to archive." Move, not copy or delete: the file is now permanent provenance (its
    # archived path is what apply_session_patch.py's _write_change_log() records as
    # record_change_log.change_source via the patch's own patch_id, escalation #836), never
    # discarded, per this session's #1 finding.
    archived_source.parent.mkdir(parents=True, exist_ok=True)
    input_path.replace(archived_source)
    return {"path": str(output), "sections": len(operations), "archived_source": str(archived_source)}


# ── flag (angle a of escalation #829 sec12.4 -- create only, no prose-section reference) ────

def run_flag(cfg, flag_code: str | None, description: str | None) -> dict:
    """Raise one `wa_data_quality_flags` instance, `flag_group='PROSE_QUALITY'`. This is the only
    direct DB write `prosestore.py` performs itself (every other operation here is read-only or
    generates a patch file for `apply_session_patch.py` to apply) — `wa_data_quality_flags` is not
    under `record_change_log` discipline (that covers `prose_section`/`prose_section_type` only,
    escalation #836), so no choke-point call applies here.

    Deliberately does NOT take a `prose_section` reference — escalation #829 sec12.2: which prose
    rows a flag touches is discovered by search at fix time (angle b, escalation #835, not built),
    not stored and kept in sync from raise time.
    """
    if not flag_code:
        raise ValueError("prose.flag needs --flag-code")
    if not description:
        raise ValueError("prose.flag needs --description")
    conn = open_db(cfg)
    try:
        row = conn.execute(
            "SELECT id FROM wa_quality_flag_types WHERE flag_group='PROSE_QUALITY' "
            "AND flag_code=? AND delete_flagged=0", (flag_code,)).fetchone()
        if not row:
            live_codes = [r["flag_code"] for r in conn.execute(
                "SELECT flag_code FROM wa_quality_flag_types "
                "WHERE flag_group='PROSE_QUALITY' AND delete_flagged=0 ORDER BY id")]
            raise ValueError(
                f"unknown --flag-code {flag_code!r} -- live PROSE_QUALITY codes: {live_codes}")
        flag_id = row["id"]
        cur = conn.execute(
            "INSERT INTO wa_data_quality_flags (flag_id, description, delete_flagged) "
            "VALUES (?,?,0)", (flag_id, description))
        conn.commit()
        return {"id": cur.lastrowid, "flag_id": flag_id, "flag_code": flag_code,
                "description": description}
    finally:
        conn.close()


# ── flag-fix (angle b of escalation #829 sec12.4 / #890 D5 -- propose -> approve -> apply) ──

def run_flag_fix_propose(cfg, flag_code: str | None, find: str | None, replace: str | None,
                          out=None) -> dict:
    """**Propose** step. Searches active `prose_section.body` for a literal `find` substring and
    writes a review report (a `.json` file, not a DB write or a patch) listing every matching
    section with its proposed replacement text. Read-only against the database — matches this
    module's existing convention (every operation except `run_flag` is read-only or
    patch-generating, never a direct write).

    Deliberately does NOT write to `record_change_log` with `status='change_proposed'`, even
    though that vocabulary already exists on the column — checked live before building (escalation
    #890): `record_change_log.payload` has its own hard rule (`record-change-log-payload-is-
    prior-state`) that it holds ONLY the content a change overwrote, never the resulting content.
    A pending proposal's payload is the opposite — the content that WOULD result, not yet applied
    — so writing it into that same field under a different status would make the field mean two
    different things depending on status, silently breaking that rule for any code/report that
    reads `payload` assuming "prior state" unconditionally. A plain review file avoids the
    collision entirely and matches how every other approval step in this project works (a written
    record the researcher reviews, then a separate explicit go-ahead) rather than inventing a new,
    narrower meaning for an already-governed field.
    """
    if not flag_code:
        raise ValueError("prose.flag_fix_propose needs --flag-code")
    if not find:
        raise ValueError("prose.flag_fix_propose needs --find")
    if replace is None:
        raise ValueError("prose.flag_fix_propose needs --replace (may be empty string)")
    conn = open_db(cfg)
    try:
        flag_row = conn.execute(
            "SELECT id FROM wa_quality_flag_types WHERE flag_group='PROSE_QUALITY' "
            "AND flag_code=? AND delete_flagged=0", (flag_code,)).fetchone()
        if not flag_row:
            live_codes = [r["flag_code"] for r in conn.execute(
                "SELECT flag_code FROM wa_quality_flag_types "
                "WHERE flag_group='PROSE_QUALITY' AND delete_flagged=0 ORDER BY id")]
            raise ValueError(
                f"unknown --flag-code {flag_code!r} -- live PROSE_QUALITY codes: {live_codes}")
        rows = conn.execute(
            """SELECT ps.id, ps.heading, ps.body, pst.code AS section_type_code
                 FROM prose_section ps
                 JOIN prose_section_type pst ON pst.id = ps.section_type_id
                WHERE COALESCE(ps.delete_flagged, 0) = 0 AND ps.body LIKE '%' || ? || '%'""",
            (find,),
        ).fetchall()
        matches = []
        for row in rows:
            body = row["body"] or ""
            matches.append({
                "prose_section_id": row["id"],
                "section_type_code": row["section_type_code"],
                "heading": row["heading"],
                "occurrences": body.count(find),
                "proposed_body": body.replace(find, replace),
            })
    finally:
        conn.close()
    report = {
        "generated_at": now_iso(), "flag_code": flag_code, "find": find, "replace": replace,
        "match_count": len(matches), "matches": matches,
    }
    stamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    output = Path(out) if out else search_output_dir(cfg) / f"prose-flag-fix-proposal-{stamp}.json"
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")
    return {"path": str(output), "match_count": len(matches)}


def run_flag_fix_apply(cfg, proposal_file, section_ids: list[int], flag_code: str | None,
                        out=None) -> dict:
    """**Apply** step (run only after the researcher has reviewed the proposal report and chosen
    which `section_ids` to act on — the "approve" stage of propose/approve/apply is that review,
    done by reading the file, not a separate mechanised gate). Re-reads each approved section's
    CURRENT body fresh from the database (never trusts the proposal file's cached snapshot, in
    case the row changed between propose and apply) and generates a `PROSE` supersede patch —
    the same, already-built-and-tested operation `run_import_chapter` already uses. Writes no DB
    row itself; apply the reviewed patch with `scripts/apply_session_patch.py`, same boundary as
    every other patch-generating operation in this module.

    Does NOT close the `wa_data_quality_flags` row itself — that write is deliberately left as a
    separate, explicit step taken once the generated patch has actually been applied (closing it
    here, before the patch is live, would let a flag read "corrected" while the body text hasn't
    actually changed yet if the patch is never applied)."""
    if not section_ids:
        raise ValueError("prose.flag_fix_apply needs --section-ids")
    proposal_path = Path(proposal_file)
    if not proposal_path.exists():
        raise FileNotFoundError(f"proposal file not found: {proposal_path}")
    proposal = json.loads(proposal_path.read_text(encoding="utf-8"))
    find = proposal["find"]
    replace = proposal["replace"]
    proposed_ids = {m["prose_section_id"] for m in proposal["matches"]}
    unknown_ids = [sid for sid in section_ids if sid not in proposed_ids]
    if unknown_ids:
        raise ValueError(
            f"section id(s) {unknown_ids} are not in the proposal file {proposal_path} -- "
            f"re-run prose.flag_fix_propose if the target set has changed")

    conn = open_db(cfg)
    operations = []
    skipped = []
    try:
        for section_id in section_ids:
            row = conn.execute(
                "SELECT id, heading, body FROM prose_section "
                "WHERE id = ? AND COALESCE(delete_flagged, 0) = 0", (section_id,)).fetchone()
            if not row:
                skipped.append({"id": section_id, "reason": "no longer an active row"})
                continue
            body = row["body"] or ""
            if find not in body:
                skipped.append({"id": section_id,
                                 "reason": f"current body no longer contains {find!r} -- content "
                                           f"changed since the proposal was generated"})
                continue
            operations.append({
                "op_id": f"PROSE-{section_id}-FLAGFIX-SUPERSEDE",
                "table": "prose_section",
                "operation": "supersede",
                "supersedes_id": section_id,
                "record": {
                    "body": body.replace(find, replace), "heading": row["heading"],
                    "author": "claude_code", "status": "draft",
                    "metadata_json": json.dumps({
                        "flag_fix": "iba.app.lib.prosestore.run_flag_fix_apply",
                        "flag_code": flag_code, "find": find, "replace": replace,
                    }),
                },
            })
    finally:
        conn.close()

    if not operations:
        raise ValueError(
            f"none of the requested section(s) are still fixable -- skipped: {skipped}")

    stamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    patch_id = f"PATCH-{stamp}-PROSE-FLAGFIX-SUPERSEDE"
    patch = {
        "_patch_meta": {
            "patch_id": patch_id, "patch_type": "PROSE", "produced_at": now_iso(),
            "session_b_status": None, "researcher_approval": "PENDING",
            "description": f"Flag-fix ({flag_code}): supersede {len(operations)} prose "
                            f"section(s), {find!r} -> {replace!r}.",
        },
        "operations": operations,
        "_patch_summary": {
            "total_operations": len(operations), "prose_section_supersedes": len(operations),
            "flag_code": flag_code, "skipped": skipped,
            "post_apply_note": "Once this patch is applied, close the corresponding "
                                "wa_data_quality_flags row(s) for this flag_code via a direct "
                                "update (corrective_action/correction_date) -- not automated by "
                                "this step, per this module's read-only/patch-generating "
                                "convention (run_flag is the sole direct-write exception).",
        },
    }
    output = Path(out) if out else patch_output_dir(cfg) / f"wa-prose-flag-fix-supersede-{stamp}.json"
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(patch, indent=2, ensure_ascii=False), encoding="utf-8")
    return {"path": str(output), "sections": len(operations), "skipped": skipped}


# ── status set/reset (escalation #918, 2026-08-27) ──────────────────────────────
def run_set_status(cfg, section_ids: list[int], status: str, author="researcher",
                    out=None) -> dict:
    """Set (or reset) `prose_section.status` for one or more sections directly, as its own
    reviewer action -- distinct from `run_import_chapter`'s content-edit round trip, which also
    happens to touch `status` (always to 'draft') as a side effect of a body rewrite. This is the
    dedicated command: read a section, decide it's read/approved (or that an earlier approval
    needs reopening), and record just that, with no body change at all.

    `status` must be a live value of `cfg_enum prose_section_status` (draft / in_review / approved
    / archived) -- checked here, not left for the applicator to discover. Generates a `PROSE`
    patch (`prose_section`/`set_status`, a new, narrower sibling of the existing `approve` op —
    that one is approve-only and one-directional; this one moves either way and is the general
    case). Writes no DB row itself; apply with `scripts/apply_session_patch.py`, the same boundary
    every other operation in this module keeps."""
    if not section_ids:
        raise ValueError("prose.set_status needs --section-ids")
    valid_statuses = cfg.enum("prose_section_status")
    if status not in valid_statuses:
        raise ValueError(f"status {status!r} is not a live value of cfg_enum "
                          f"'prose_section_status' ({valid_statuses!r})")

    conn = open_db(cfg)
    operations = []
    skipped = []
    try:
        for section_id in section_ids:
            row = conn.execute(
                "SELECT id, status FROM prose_section "
                "WHERE id = ? AND COALESCE(delete_flagged, 0) = 0", (section_id,)).fetchone()
            if not row:
                skipped.append({"id": section_id, "reason": "no longer an active row"})
                continue
            if row["status"] == status:
                skipped.append({"id": section_id, "reason": f"already {status!r} -- no-op"})
                continue
            operations.append({
                "op_id": f"PROSE-{section_id}-SET-STATUS",
                "table": "prose_section",
                "operation": "set_status",
                "id": section_id,
                "record": {"status": status, "approved_by": author if status == "approved" else None},
            })
    finally:
        conn.close()

    if not operations:
        raise ValueError(f"no section needed a status change -- skipped: {skipped}")

    stamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    patch_id = f"PATCH-{stamp}-PROSE-SET-STATUS"
    patch = {
        "_patch_meta": {
            "patch_id": patch_id, "patch_type": "PROSE", "produced_at": now_iso(),
            "session_b_status": None, "researcher_approval": "PENDING",
            "description": f"Set {len(operations)} prose_section row(s) to status={status!r}.",
        },
        "operations": operations,
        "_patch_summary": {
            "total_operations": len(operations), "prose_section_status_set": len(operations),
            "status": status, "skipped": skipped,
        },
    }
    output = Path(out) if out else patch_output_dir(cfg) / f"wa-prose-set-status-{stamp}.json"
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(patch, indent=2, ensure_ascii=False), encoding="utf-8")
    return {"path": str(output), "sections": len(operations), "skipped": skipped}
